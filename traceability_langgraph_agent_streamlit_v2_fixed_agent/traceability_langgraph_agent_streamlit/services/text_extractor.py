from pathlib import Path
import pandas as pd
from docx import Document
import fitz

def read_text_safely(path: str) -> str:
    for enc in ["utf-8", "utf-8-sig", "cp1252", "latin1"]:
        try:
            return Path(path).read_text(encoding=enc)
        except UnicodeDecodeError:
            pass
    return Path(path).read_text(encoding="latin1", errors="replace")

def extract_pdf(path: str) -> dict:
    doc = fitz.open(path)
    chunks, full = [], ""
    for i, page in enumerate(doc, start=1):
        t = page.get_text() or ""
        chunks.append({"source_ref": f"Page {i}", "text": t})
        full += f"\n--- Page {i} ---\n{t}"
    doc.close()
    return {"full_text": full, "chunks": chunks}

def extract_docx(path: str) -> dict:
    doc = Document(path)
    blocks = []
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            rt = " | ".join(c.text.strip() for c in row.cells)
            if rt.strip():
                blocks.append(rt)
    return {"full_text": "\n".join(blocks), "chunks": [{"source_ref": f"Block {i+1}", "text": b} for i,b in enumerate(blocks)]}

def extract_excel(path: str) -> dict:
    sheets = pd.read_excel(path, sheet_name=None)
    full, chunks = "", []
    for sheet, df in sheets.items():
        df = df.fillna("").astype(str)
        full += f"\n--- Sheet: {sheet} ---\n" + df.to_string(index=False)
        for idx, row in df.iterrows():
            rt = " | ".join([str(x) for x in row.tolist() if str(x).strip()])
            if rt.strip():
                chunks.append({"source_ref": f"Sheet {sheet}, Row {idx+2}", "text": rt})
    return {"full_text": full, "chunks": chunks}

def extract_csv(path: str) -> dict:
    last = None
    for enc in ["utf-8", "utf-8-sig", "cp1252", "latin1"]:
        try:
            df = pd.read_csv(path, encoding=enc).fillna("").astype(str)
            full = df.to_string(index=False)
            chunks = []
            for idx, row in df.iterrows():
                rt = " | ".join([str(x) for x in row.tolist() if str(x).strip()])
                if rt.strip():
                    chunks.append({"source_ref": f"Row {idx+2}", "text": rt})
            return {"full_text": full, "chunks": chunks}
        except Exception as e:
            last = e
    raw = read_text_safely(path)
    return {"full_text": raw, "chunks": [{"source_ref": "Raw Text", "text": raw}], "warning": str(last)}

def extract_txt(path: str) -> dict:
    raw = read_text_safely(path)
    return {"full_text": raw, "chunks": [{"source_ref": f"Line {i+1}", "text": line} for i,line in enumerate(raw.splitlines()) if line.strip()]}

def extract_full_text(path: str) -> dict:
    ext = Path(path).suffix.lower()
    if ext == ".pdf": return extract_pdf(path)
    if ext == ".docx": return extract_docx(path)
    if ext in [".xlsx", ".xls"]: return extract_excel(path)
    if ext == ".csv": return extract_csv(path)
    if ext == ".txt": return extract_txt(path)
    raise ValueError(f"Unsupported file type: {ext}")

def get_content_summary(text: str) -> dict:
    return {"word_count": len(text.split()), "character_count": len(text), "preview": text[:2500]}
